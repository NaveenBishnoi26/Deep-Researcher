"""
Report export module for the Deep Researcher application.

This module handles the export of research reports in various formats:
1. PDF
2. DOCX
3. HTML
"""

import os
from typing import Dict, Any, List
import logging
from pathlib import Path
from docx import Document as DocxDocument
import markdown
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Frame, PageTemplate, BaseDocTemplate
from reportlab.lib.units import inch
import shutil
from datetime import datetime
import re
from reportlab.pdfgen.canvas import Canvas
from reportlab.platypus.tableofcontents import TableOfContents
from reportlab.platypus.flowables import Flowable
import pdfkit

from backend.config import (
    EXPORTS_DIR,
    PDF_OPTIONS
)

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Define a page number class for PDF
class PageNumCanvas(Canvas):
    """Canvas with page numbers."""
    
    def __init__(self, *args, **kwargs):
        """Initialize the canvas."""
        Canvas.__init__(self, *args, **kwargs)
        self.pages = []
        
    def showPage(self):
        """Add page info to the list of pages."""
        self.pages.append(dict(self.__dict__))
        self._startPage()
        
    def save(self):
        """Add page numbers to the document."""
        page_count = len(self.pages)
        
        for page in self.pages:
            self.__dict__.update(page)
            self.draw_page_number(page_count)
            Canvas.showPage(self)
            
        Canvas.save(self)
        
    def draw_page_number(self, page_count):
        """Draw the page number on the canvas."""
        page = f"Page {self._pageNumber} of {page_count}"
        self.setFont("Helvetica", 9)
        self.drawRightString(letter[0] - 50, 40, page)


class PDFReport(BaseDocTemplate):
    """Custom PDF document with page numbers and headers."""
    
    def __init__(self, filename, title="", **kwargs):
        """Initialize the PDF document."""
        self.title = title
        self.page_info = title
        
        BaseDocTemplate.__init__(self, filename, **kwargs)
        
        # Define frames
        frame = Frame(
            self.leftMargin, 
            self.bottomMargin, 
            self.width, 
            self.height - 50,  # Leave space for header
            id='normal'
        )
        
        # Define page templates
        template = PageTemplate(
            id='default',
            frames=[frame],
            onPage=self._header_footer
        )
        
        self.addPageTemplates([template])
        
    def _header_footer(self, canvas, doc):
        """Add header and footer to each page."""
        # Save state
        canvas.saveState()
        
        # Header - title
        canvas.setFont('Helvetica-Bold', 10)
        canvas.drawString(doc.leftMargin, doc.height + doc.topMargin - 20, self.title)
        
        # Header - line
        canvas.line(doc.leftMargin, doc.height + doc.topMargin - 30, 
                   doc.width + doc.leftMargin, doc.height + doc.topMargin - 30)
        
        # Footer - page number
        canvas.setFont('Helvetica', 9)
        page_num = f"Page {doc.page} of {doc._pageCount}"
        canvas.drawRightString(doc.width + doc.leftMargin, doc.bottomMargin - 20, page_num)
        
        # Footer - timestamp
        timestamp = datetime.now().strftime('%Y-%m-%d')
        canvas.drawString(doc.leftMargin, doc.bottomMargin - 20, f"Generated: {timestamp}")
        
        # Restore state
        canvas.restoreState()

class ReportExporter:
    """Handles the export of research reports in various formats."""
    
    def __init__(self):
        """Initialize the report exporter."""
        self.exports_dir = Path(EXPORTS_DIR)
        self.exports_dir.mkdir(parents=True, exist_ok=True)
        
        # Define styles
        self.styles = getSampleStyleSheet()
        self.styles.add(ParagraphStyle(
            name='CustomCitation',
            parent=self.styles['Normal'],
            fontSize=10,
            leftIndent=20,
            spaceBefore=6,
            spaceAfter=6
        ))
        self.styles.add(ParagraphStyle(
            name='CustomHeading1',
            parent=self.styles['Heading1'],
            fontSize=16,
            spaceBefore=24,
            spaceAfter=12
        ))
        self.styles.add(ParagraphStyle(
            name='CustomHeading2',
            parent=self.styles['Heading2'],
            fontSize=14,
            spaceBefore=18,
            spaceAfter=8
        ))
        self.styles.add(ParagraphStyle(
            name='CustomNormal',
            parent=self.styles['Normal'],
            fontSize=12,
            spaceBefore=6,
            spaceAfter=6
        ))
    
    def export_pdf(self, content: str, citations: List[Dict[str, str]], filename: str) -> str:
        """Export report to PDF format."""
        try:
            # Configure PDF options
            options = {
                'page-size': 'A4',
                'margin-top': '20mm',
                'margin-right': '20mm',
                'margin-bottom': '20mm',
                'margin-left': '20mm',
                'encoding': 'UTF-8',
                'custom-header': [
                    ('Accept-Encoding', 'gzip')
                ],
                'no-outline': None
            }
            
            # Add custom CSS for improved formatting mimicking Gemini sample
            css = """
            body        { font-family: "Georgia", serif; font-size: 11pt; line-height: 1.45; }
            h1          { font-size: 20pt; margin-top: 24pt; margin-bottom: 12pt; }
            h2          { font-size: 16pt; margin-top: 18pt; margin-bottom:  9pt; }
            h3          { font-size: 14pt; margin-top: 14pt; margin-bottom:  7pt; }
            p, li       { margin: 6pt 0 6pt 0; }
            .toc h1     { font-size: 14pt; margin: 0 0 8pt 0; }
            .toc ul li  { list-style-type: none; margin-left: 0; }
            """
            
            # NEW: convert markdown to HTML first
            md_html = markdown.markdown(
                content,
                extensions=[
                    "fenced_code",
                    "tables",
                    "toc",            # auto-generates <div class="toc">…</div>
                    "nl2br",
                    "sane_lists"
                ]
            )
            
            # Extract TOC if present and reposition it
            toc_html = re.search(r'<div class="toc">(.*?)</div>', md_html, flags=re.DOTALL)
            toc_block = toc_html.group(0) if toc_html else ""
            md_html = md_html.replace(toc_block, "")  # strip from body; we'll re-insert manually
            
            # Create HTML content
            html_content = f"""
            <html>
            <head>
                <style>{css}</style>
            </head>
            <body>
                {toc_block}
                {md_html}
                <h2>References</h2>
                <ol>
                    {''.join(f'<li>{citation.get("text", "")}</li>' for citation in citations)}
                </ol>
            </body>
            </html>
            """
            
            # Get wkhtmltopdf path
            wkhtmltopdf_path = r'C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe'
            if not os.path.exists(wkhtmltopdf_path):
                wkhtmltopdf_path = None  # Let pdfkit find it automatically
            
            # Configure pdfkit
            config = pdfkit.configuration(wkhtmltopdf=wkhtmltopdf_path)
            
            # Generate PDF
            output_path = os.path.join(EXPORTS_DIR, f"{filename}.pdf")
            pdfkit.from_string(html_content, output_path, options=options, configuration=config)
            
            return output_path
            
        except Exception as e:
            logger.error(f"Error exporting PDF: {str(e)}")
            raise

    def _export_error_pdf(self, filename: str) -> str:
        """Generate a very basic error PDF as last resort."""
        error_filename = f"{filename}_error.pdf" if filename else f"error_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        output_path = self.exports_dir / error_filename
        
        # Create a very simple PDF with minimal content
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        # Minimal error content
        content = [
            Paragraph("Error Generating Report", self.styles["Title"]),
            Spacer(1, 24),
            Paragraph("There was an error generating the PDF report. Please try again or contact support.", self.styles["Normal"]),
            Spacer(1, 12),
            Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", self.styles["Normal"])
        ]
        
        # Build the error PDF
        doc.build(content)
        logger.info(f"Generated error PDF: {output_path}")
        return str(output_path)
        
    def _create_minimal_pdf(self, filename: str) -> str:
        """
        Create a minimal PDF with an error message as a last resort.
        
        Args:
            filename: The filename for the PDF
            
        Returns:
            The path to the created PDF file
        """
        try:
            # Ensure the filename has no extension (we'll add .pdf)
            filename = os.path.splitext(filename)[0]
            
            # Create PDF document in exports directory
            output_path = self.exports_dir / f"{filename}_minimal.pdf"
            
            # Make sure the exports directory exists
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # Create a very simple PDF with minimal content
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            
            # Use the simplest possible PDF creation method
            c = canvas.Canvas(str(output_path), pagesize=letter)
            width, height = letter
            
            # Add title
            c.setFont("Helvetica-Bold", 16)
            c.drawString(72, height - 72, "Research Report")
            
            # Add timestamp
            c.setFont("Helvetica", 10)
            timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            c.drawString(72, height - 100, f"Generated on: {timestamp}")
            
            # Add error message
            c.setFont("Helvetica", 12)
            c.drawString(72, height - 150, "An error occurred during PDF generation.")
            c.drawString(72, height - 170, "Please try regenerating the report.")
            
            # Save the PDF
            c.save()
            
            logger.info(f"Created minimal PDF at {output_path}")
            return str(output_path)
            
        except Exception as e:
            # Last resort: create an absolutely minimal PDF
            logger.error(f"Error creating minimal PDF: {str(e)}")
            try:
                # Use an even more minimal approach without any external modules
                minimal_path = self.exports_dir / f"{filename}_fallback.pdf"
                
                # Create a bare-minimum valid PDF
                with open(minimal_path, 'wb') as f:
                    f.write(b"%PDF-1.4\n")
                    f.write(b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n")
                    f.write(b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n")
                    f.write(b"3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R/Resources<<>>/Contents 4 0 R>>endobj\n")
                    f.write(b"4 0 obj<</Length 21>>stream\nBT /F1 12 Tf 100 700 Td (Research Report Error) Tj ET\nendstream\nendobj\n")
                    f.write(b"xref\n0 5\n0000000000 65535 f\n0000000010 00000 n\n0000000053 00000 n\n0000000102 00000 n\n0000000192 00000 n\ntrailer<</Size 5/Root 1 0 R>>\nstartxref\n270\n%%EOF")
                
                logger.info(f"Created fallback PDF at {minimal_path}")
                return str(minimal_path)
            except Exception as final_error:
                logger.error(f"Final fallback PDF creation failed: {str(final_error)}")
                # Return the path anyway, even though the file might not exist or be valid
                return str(self.exports_dir / f"{filename}_error.pdf")
    
    def export_docx(self, report: Dict[str, Any], filename: str) -> str:
        """
        Export the report to DOCX format.
        
        Args:
            report: The report dictionary
            filename: Base filename for the export
            
        Returns:
            Path to the exported DOCX file
        """
        try:
            # Create a new document
            doc = DocxDocument()
            
            # Add title
            doc.add_heading(report.get('title', 'Untitled Report'), 0)
            
            # Add date/timestamp
            timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            doc.add_paragraph(f"Generated on: {timestamp}")
            
            # Check if there are any sections
            if not report.get('sections'):
                doc.add_heading('Report Content', 1)
                doc.add_paragraph("Report content is being generated. Please check back later for more detailed information.")
                # Save the document
                output_path = os.path.join(EXPORTS_DIR, f"{filename}.docx")
                doc.save(output_path)
                return output_path
            
            # Add sections
            for section in report.get('sections', []):
                # Add section title
                doc.add_heading(section.get('section_title', 'Untitled Section'), 1)
                
                # Add section content
                if section.get('content') and section['content']:
                    doc.add_paragraph(section['content'])
                else:
                    doc.add_paragraph("This section is being generated. More information will be available soon.")
                
                # Add subsections
                if section.get('subsections'):
                    for subsection in section['subsections']:
                        # Add subsection title
                        doc.add_heading(subsection.get('subsection_title', 'Untitled Subsection'), 2)
                        
                        # Add subsection content
                        if subsection.get('content') and subsection['content']:
                            doc.add_paragraph(subsection['content'])
                        else:
                            doc.add_paragraph("This subsection is being generated. More information will be available soon.")
            
            # Add glossary if it exists
            if report.get('glossary'):
                doc.add_heading('Glossary', 1)
                for entry in report['glossary']:
                    doc.add_paragraph(f"{entry['term']}: {entry['definition']}")
            
            # Add references if available
            if report.get('references'):
                doc.add_heading('References', 1)
                for i, citation in enumerate(report.get('references', []), 1):
                    doc.add_paragraph(f"{i}. {citation}")
            
            # Save the document
            output_path = os.path.join(EXPORTS_DIR, f"{filename}.docx")
            doc.save(output_path)
            
            logger.info(f"Exported DOCX to {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error exporting DOCX: {str(e)}")
            raise
    
    def export_html(self, report: Dict[str, Any], filename: str) -> str:
        """
        Export the report as an HTML file.
        
        Args:
            report: Dictionary containing the report content
            filename: Name of the output file
            
        Returns:
            Path to the exported file
        """
        try:
            # Format report content in markdown format
            md_content = self._format_report_content(report)
            
            # Convert markdown to HTML
            html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{report.get('title', 'Research Report')}</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            color: #333;
        }}
        h1 {{
            color: #1a1a1a;
            border-bottom: 2px solid #ddd;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #2a2a2a;
            margin-top: 30px;
        }}
        h3 {{
            color: #3a3a3a;
            margin-top: 20px;
        }}
        p {{
            margin-bottom: 16px;
        }}
        .references {{
            margin-top: 40px;
            border-top: 1px solid #ddd;
            padding-top: 20px;
        }}
        .reference-item {{
            padding: 5px 0;
        }}
        .timestamp {{
            color: #666;
            font-style: italic;
            margin-bottom: 30px;
        }}
    </style>
</head>
<body>
    <h1>{report.get('title', 'Research Report')}</h1>
    <div class="timestamp">Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</div>
    
    {markdown.markdown(md_content)}
    
    <div class="references">
        <h2>References</h2>
        <ol>
"""
            
            # Add citations
            if report.get("citations"):
                for citation in report.get("citations"):
                    if isinstance(citation, dict):
                        if 'text' in citation:
                            citation_text = citation['text']
                        else:
                            citation_text = citation.get('title', 'Untitled')
                            if citation.get("url"):
                                citation_text += f" - {citation['url']}"
                            if citation.get("authors"):
                                citation_text += f" by {citation['authors']}"
                    else:
                        citation_text = str(citation)
                        
                    html_content += f"            <li class=\"reference-item\">{citation_text}</li>\n"
            else:
                html_content += "            <li>No citations available</li>\n"
                
            html_content += """        </ol>
    </div>
</body>
</html>"""
            
            # Save the HTML file
            output_path = os.path.join(EXPORTS_DIR, f"{filename}.html")
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(html_content)
            
            logger.info(f"Exported HTML to {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error exporting HTML: {str(e)}")
            raise

    def _format_report_content(self, report: Dict[str, Any]) -> str:
        """Format report content for export."""
        content = []
        
        # Add title
        content.append(f"# {report['title']}\n")
        
        # Add date
        if 'date' in report:
            content.append(f"*Generated on: {report['date']}*\n")
        
        # Check if there are any sections
        if not report.get('sections'):
            content.append("\n## Report Content\n")
            content.append("Report content is being generated. Please check back later for more detailed information.\n")
            return "\n".join(content)
        
        # Process sections
        for section in report['sections']:
            # Add section title
            content.append(f"\n## {section['section_title']}\n")
            
            # Add section content
            if section.get('content'):
                content.append(f"{section['content']}\n")
            else:
                content.append("This section is being generated. More information will be available soon.\n")
            
            # Process subsections
            if section.get('subsections'):
                for subsection in section['subsections']:
                    # Add subsection title
                    content.append(f"\n### {subsection['subsection_title']}\n")
                    
                    # Add subsection content
                    if subsection.get('content'):
                        content.append(f"{subsection['content']}\n")
                    else:
                        content.append("This subsection is being generated. More information will be available soon.\n")
        
        # Add conclusion if available
        if report.get('conclusion'):
            content.append("\n## Conclusion\n")
            content.append(f"{report['conclusion']}\n")
        
        # Add references if available
        if report.get('references'):
            content.append("\n## References\n")
            for i, citation in enumerate(report['references'], 1):
                content.append(f"{i}. {citation}\n")
        
        return "\n".join(content)

# Example usage
if __name__ == "__main__":
    exporter = ReportExporter()
    
    # Example report
    report = {
        "title": "Test Report",
        "content": "# Test Report\n\nThis is a test report.",
        "citations": [
            "Author, A. (2023). Test Paper. Journal of Testing, 1(1), 1-10."
        ]
    }
    
    # Export in different formats
    pdf_path = exporter.export_pdf(report["content"], report["citations"], "test_report")
    docx_path = exporter.export_docx(report, "test_report")
    html_path = exporter.export_html(report, "test_report")
    
    print(f"Exported to:\nPDF: {pdf_path}\nDOCX: {docx_path}\nHTML: {html_path}") 